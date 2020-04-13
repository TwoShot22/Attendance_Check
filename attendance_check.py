import docx2txt
import docx
import re

########################################################################
# 채팅 Docx 파일의 Data 분석

# 학생 명단 Data 불러오기
student_docx = docx2txt.process("Document/학생_명단.docx")

# Meet의 채팅 기록 Data 불러오기
chat_docx = docx2txt.process("Document/1주차_출석.docx")

# 문자열로 저장된 Data를 단어 단위로 배열로 저장
student_array = student_docx.split()
chat_array = chat_docx.split()

# 학생 : 언급 된 횟수를 저장하기 위한 Dictionary
attendance_data = dict()

# 학생 별 언급 횟수를 Count하기 위한 Function
for i in range(len(student_array)):
    check_student = student_array[i]
    check_number = chat_array.count("그래픽아츠과/학생/"+check_student)
    
    attendance_data[check_student] = check_number

print("--- Data 분석 완료 ---")
########################################################################
# 분석된 Data를 Docx로 저장

# Docx 생성
doc = docx.Document()
title = "학생 이름 / 언급 횟수"
doc_head = doc.add_paragraph()
doc_head.add_run(title).bold = True

# Docx에 Data 추가
for key, value in attendance_data.items():
    result_student = str(key)
    result_attendant = str(value)

    result = result_student + "       /       " + result_attendant
    doc.add_paragraph(result)

# Docx 저장
doc.save('Document/출석_체크.docx')
print("--- Docx 저장 완료 ---")